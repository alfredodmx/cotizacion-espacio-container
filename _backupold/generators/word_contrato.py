"""
Generador de contrato Word (.docx) y helpers de conversion de montos a palabras.
"""


def num_a_palabras(n: int) -> str:
    """Convierte un numero entero a su representacion en palabras en espanol."""
    unidades = ['','uno','dos','tres','cuatro','cinco','seis','siete','ocho','nueve',
                'diez','once','doce','trece','catorce','quince','dieciséis','diecisiete',
                'dieciocho','diecinueve','veinte','veintiuno','veintidós','veintitrés',
                'veinticuatro','veinticinco','veintiséis','veintisiete','veintiocho','veintinueve']
    decenas  = ['','diez','veinte','treinta','cuarenta','cincuenta','sesenta','setenta','ochenta','noventa']
    centenas = ['','ciento','doscientos','trescientos','cuatrocientos','quinientos',
                'seiscientos','setecientos','ochocientos','novecientos']
    if n == 0: return 'cero'
    if n < 0:  return 'menos ' + num_a_palabras(-n)
    res = ''
    if n >= 1_000_000:
        m = n // 1_000_000
        res += ('un millón' if m == 1 else num_a_palabras(m) + ' millones') + ' '
        n %= 1_000_000
    if n >= 1000:
        m = n // 1000
        res += ('mil' if m == 1 else num_a_palabras(m) + ' mil') + ' '
        n %= 1000
    if n >= 100:
        if n == 100: res += 'cien '
        else: res += centenas[n // 100] + ' '
        n %= 100
    if n >= 30:
        d, u = divmod(n, 10)
        res += decenas[d] + (' y ' + unidades[u] if u else '') + ' '
    elif n > 0:
        res += unidades[n] + ' '
    return res.strip()


def monto_a_palabras(monto: float) -> str:
    """Convierte monto a texto: 'X pesos'."""
    return num_a_palabras(int(round(monto))) + ' pesos'


def generar_word_contrato(datos: dict):
    """
    Genera un .docx del contrato con los datos ya rellenados.
    Los campos autollenados aparecen en AZUL NEGRITA.
    Retorna un BytesIO con el .docx.
    """
    try:
        from docx import Document as _DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io as _io
    except ImportError:
        return None

    doc = _DocxDocument()

    for section in doc.sections:
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)
        section.top_margin    = Cm(3.8)
        section.bottom_margin = Cm(2.2)

    AZUL = RGBColor(0x0f, 0x34, 0x60)

    def _style_normal(paragraph):
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _run_normal(paragraph, text, bold=False, color=None, size=10.5):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        if color:
            run.font.color.rgb = color
        return run

    def _run_dato(paragraph, text):
        return _run_normal(paragraph, text, bold=True, color=AZUL)

    def _seccion(doc, texto):
        p = doc.add_paragraph()
        run = p.add_run(f'  {texto}  ')
        run.bold = True; run.font.size = Pt(9.5)
        run.font.name = 'Arial'; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F3460')
        pPr.append(shd)
        return p

    def _hr(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'E8EEF7')
        pBdr.append(bottom); pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(6)

    d = datos
    fmt = lambda v: "${:,.0f}".format(v).replace(",", ".")

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = nota.add_run('CAMPOS EN AZUL NEGRITA = AUTOLLENADOS — NO MODIFICAR')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    nota.paragraph_format.space_after = Pt(12)

    t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t1.add_run('CONTRATO DE FABRICACIÓN Y VENTA')
    r.bold = True; r.font.size = Pt(15); r.font.name = 'Times New Roman'; r.font.color.rgb = AZUL
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run('VIVIENDA TIPO CONTAINER')
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'; r.font.color.rgb = AZUL

    doc.add_paragraph()

    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'En Santiago de Chile, a ')
    _run_dato(p, d['fecha_str'])
    _run_normal(p, ', comparecen:')

    _seccion(doc, 'I. COMPARECENCIA')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, '1. EL PROVEEDOR', bold=True)

    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'Don ')
    _run_normal(p, 'Alan Mauricio Gatica Concha', bold=True)
    _run_normal(p, ', cédula nacional de identidad N° ')
    _run_normal(p, '13.668.157-5', bold=True)
    _run_normal(p, ', en representación de ')
    _run_normal(p, 'Inversiones Container House SpA', bold=True)
    _run_normal(p, ', Rol Único Tributario N° ')
    _run_normal(p, '78.268.851-0', bold=True)
    _run_normal(p, ', ambos con domicilio para estos efectos en Villasana N° 2039, Departamento 51, Torre D, comuna de Quinta Normal, Región Metropolitana, quien en adelante se denominará indistintamente "el Proveedor".')

    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, '2. EL CLIENTE', bold=True)

    p = doc.add_paragraph(); _style_normal(p)
    tratamiento = d.get('cli_tratamiento', 'Don')
    _run_normal(p, f'{tratamiento} ')
    _run_dato(p, d['cli_nombre'])
    _run_normal(p, ', cédula nacional de identidad N° ')
    _run_dato(p, d['cli_rut'])
    if d.get('tipo_cliente') == 'juridica':
        _run_normal(p, ', en representación de ')
        _run_dato(p, d.get('cli_empresa', ''))
        _run_normal(p, ', RUT ')
        _run_dato(p, d.get('cli_rut_empresa', ''))
    _run_normal(p, ', con domicilio en ')
    _run_dato(p, d['cli_domicilio'])
    _run_normal(p, ', comuna de ')
    _run_dato(p, d['cli_comuna'])
    _run_normal(p, f", Región {d['cli_region']}, quien en adelante se denominará \"el Cliente\".")

    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'Se deja expresa constancia que la dirección de instalación del proyecto será ')
    _run_dato(p, d['inst_domicilio'])
    _run_normal(p, ', comuna de ')
    _run_dato(p, d['inst_comuna'])
    _run_normal(p, f", Región {d['inst_region']}.")

    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'Las partes declaran ser mayores de edad, con plena capacidad legal para contratar, y acuerdan celebrar el presente ')
    _run_normal(p, 'Contrato de Fabricación y Venta de Vivienda Tipo Container', bold=True)
    _run_normal(p, ', el cual se regirá por las cláusulas que se indican a continuación.')
    _hr(doc)

    _seccion(doc, 'II. DEFINICIONES')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'Para efectos del presente contrato, se entenderá por:')

    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'a) '); _run_normal(p, 'Proyecto', bold=True)
    _run_normal(p, ': La vivienda tipo container identificada como ')
    _run_normal(p, 'Proyecto N° ', bold=True)
    _run_dato(p, d['ep_numero'])
    _run_normal(p, ' – "'); _run_dato(p, d['ep_nombre']); _run_normal(p, '".')

    for txt in [
        'b) <Anexos>: Los documentos técnicos y comerciales que forman parte integrante del presente contrato, en especial Anexo N°1 (Especificaciones Técnicas) y Anexo N°2 (Presupuesto Detallado).',
        'c) <Preentrega>: Instancia de revisión visual del módulo previo a su despacho desde las instalaciones del Proveedor.',
    ]:
        p = doc.add_paragraph(); _style_normal(p); _run_normal(p, txt)
    _hr(doc)

    _seccion(doc, 'III. OBJETO DEL CONTRATO')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'El Cliente encarga al Proveedor la ')
    _run_normal(p, 'fabricación y venta', bold=True)
    _run_normal(p, ' del Proyecto individualizado precedentemente, conforme a los ')
    _run_normal(p, 'planos entregados por el Cliente', bold=True)
    _run_normal(p, ', a las ')
    _run_normal(p, 'especificaciones técnicas', bold=True)
    _run_normal(p, ', y al ')
    _run_normal(p, 'presupuesto detallado contenido en el Anexo N°2', bold=True)
    _run_normal(p, ', documentos que el Cliente declara conocer, aceptar y que forman parte integrante e inseparable del presente contrato.')
    _hr(doc)

    _seccion(doc, 'VI. PRECIO')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'El precio total del Proyecto asciende a la suma de ')
    _run_dato(p, fmt(d["precio_total"]))
    _run_normal(p, f' ({monto_a_palabras(d["precio_total"])}), IVA incluido.')
    _hr(doc)

    _seccion(doc, 'VII. FORMA Y ETAPAS DE PAGO')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'a) '); _run_normal(p, '50% inicial', bold=True); _run_normal(p, ': ')
    _run_dato(p, fmt(d["pago_50"]))
    _run_normal(p, f' ({monto_a_palabras(d["pago_50"])}), asignación del contenedor y ejecución de obra gruesa.')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'b) '); _run_normal(p, '25% intermedio', bold=True); _run_normal(p, ': ')
    _run_dato(p, fmt(d["pago_25a"]))
    _run_normal(p, f' ({monto_a_palabras(d["pago_25a"])}), una vez finalizada la obra gruesa.')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'c) '); _run_normal(p, '25% final', bold=True); _run_normal(p, ': ')
    _run_dato(p, fmt(d["pago_25b"]))
    _run_normal(p, f' ({monto_a_palabras(d["pago_25b"])}), luego de la preentrega del Proyecto y el mismo día del despacho del módulo.')
    _hr(doc)

    _seccion(doc, 'X. PLAZO DE FABRICACIÓN Y ENTREGA')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'El plazo máximo de fabricación y entrega será de ')
    _run_dato(p, f"{d['plazo_dias']} días hábiles administrativos")
    _run_normal(p, ', contados desde el día hábil siguiente a aquel en que los fondos del anticipo se encuentren efectivamente liberados.')
    _hr(doc)

    doc.add_page_break()
    _seccion(doc, 'XVI. FIRMA')
    p = doc.add_paragraph(); _style_normal(p)
    _run_normal(p, 'El presente contrato se firma en ')
    _run_normal(p, 'dos ejemplares de igual tenor y fecha', bold=True)
    _run_normal(p, ', quedando uno en poder de cada parte.')

    doc.add_paragraph(); doc.add_paragraph()

    firma_tbl = doc.add_table(rows=4, cols=2)
    firma_tbl.cell(0, 0).paragraphs[0].add_run('EL PROVEEDOR').bold = True
    firma_tbl.cell(0, 1).paragraphs[0].add_run('EL CLIENTE').bold = True
    firma_tbl.cell(1, 0).paragraphs[0].add_run('_' * 30)
    firma_tbl.cell(1, 1).paragraphs[0].add_run('_' * 30)
    firma_tbl.cell(2, 0).paragraphs[0].add_run('Alan Mauricio Gatica Concha').bold = True
    r2 = firma_tbl.cell(2, 1).paragraphs[0].add_run(d['cli_nombre'])
    r2.bold = True; r2.font.color.rgb = AZUL
    firma_tbl.cell(3, 0).paragraphs[0].add_run('RUT: 13.668.157-5')
    r3 = firma_tbl.cell(3, 1).paragraphs[0].add_run(f"RUT: {d['cli_rut']}")
    r3.font.color.rgb = AZUL
    for row in firma_tbl.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
